import os
import json
import time
#import fitz
import pdfplumber
from docx import Document
from langchain_openai import OpenAIEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re
import logging
from typing import List, Dict, Tuple, Any

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download NLTK data if needed
import nltk
try:
    from nltk.tokenize import sent_tokenize
except LookupError:
    logger.info("Downloading NLTK punkt tokenizer...")
    nltk.download('punkt')
    from nltk.tokenize import sent_tokenize

# Updated OpenAI import
from openai import OpenAI

# -----------------------------
# API Keys & Client
# -----------------------------
from dotenv import load_dotenv
load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# -----------------------------
# 1. Parse PDF or DOCX with Error Handling
# -----------------------------


def parse_pdf(path: str) -> List[Dict[str, Any]]:
    """Parse PDF and extract text with page numbers using pdfplumber."""
    try:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    texts.append({
                        "page": page_num, 
                        "section": "Unknown", 
                        "text": text.strip()
                    })
        
        logger.info(f"Successfully parsed PDF: {len(texts)} pages")
        return texts
    except Exception as e:
        logger.error(f"Error parsing PDF {path}: {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")

def parse_docx(path: str) -> List[Dict[str, Any]]:
    """Parse DOCX and extract text with section headings."""
    try:
        doc = Document(path)
        texts = []
        current_section = "Unknown"
        page_counter = 1
        
        for para in doc.paragraphs:
            if para.style and "Heading" in para.style.name:
                current_section = para.text.strip() if para.text.strip() else "Unknown"
            elif para.text.strip():
                texts.append({
                    "page": page_counter, 
                    "section": current_section, 
                    "text": para.text.strip()
                })
                # Approximate page breaks (every 20 paragraphs)
                if len(texts) % 20 == 0:
                    page_counter += 1
        
        logger.info(f"Successfully parsed DOCX: {len(texts)} paragraphs")
        return texts
    except Exception as e:
        logger.error(f"Error parsing DOCX {path}: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def load_document(path: str) -> List[Dict[str, Any]]:
    """Load document with proper error handling."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    if path.lower().endswith(".pdf"):
        return parse_pdf(path)
    elif path.lower().endswith(".docx"):
        return parse_docx(path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

# -----------------------------
# 2. Enhanced Chunking with Classification
# -----------------------------
def classify_chunk_type(text: str) -> str:
    """Classify chunks by content type to improve retrieval precision."""
    text_lower = text.lower()
    
    if any(word in text_lower for word in ["will be", "shall", "is planned", "procedure", "visit", "assessment"]):
        return "planned_action"
    elif any(word in text_lower for word in ["was", "were", "has been", "completed", "previous", "baseline"]):
        return "completed_action"  
    elif any(word in text_lower for word in ["risk", "adverse", "side effect", "toxicity", "harm", "danger"]):
        return "risk_info"
    elif any(word in text_lower for word in ["benefit", "advantage", "improvement", "positive", "helpful"]):
        return "benefit_info"
    elif any(word in text_lower for word in ["purpose", "objective", "aim", "goal", "hypothesis", "primary endpoint"]):
        return "purpose_info"
    elif any(word in text_lower for word in ["inclusion", "exclusion", "eligibility", "criteria"]):
        return "eligibility_info"
    else:
        return "general"

def chunk_texts(texts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Chunk texts while preserving metadata and adding classification."""
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", ";", " "]
        )
        chunks = []
        for t in texts:
            text_chunks = splitter.split_text(t["text"])
            for chunk in text_chunks:
                if chunk.strip():  # Only add non-empty chunks
                    chunk_type = classify_chunk_type(chunk)
                    chunks.append({
                        "text": chunk.strip(),
                        "section": t["section"],
                        "page": t["page"],
                        "chunk_type": chunk_type
                    })
        
        logger.info(f"Created {len(chunks)} classified chunks from {len(texts)} text segments")
        return chunks
    except Exception as e:
        logger.error(f"Error chunking texts: {str(e)}")
        raise ValueError(f"Failed to chunk texts: {str(e)}")

# -----------------------------
# 3. Enhanced Retriever with Context Awareness
# -----------------------------
def build_retriever(chunks: List[Dict[str, Any]], index_name: str = "icf-protocol-index", 
                   namespace: str = "default") -> BM25Retriever:
    """Build enhanced BM25 retriever with context-aware chunking."""
    try:
        texts = [c["text"] for c in chunks]
        metadatas = [{"section": c["section"], "page": c["page"], "chunk_type": c["chunk_type"]} 
                    for c in chunks]

        if not texts:
            raise ValueError("No text chunks provided for retriever")

        # Build enhanced BM25 retriever
        retriever = BM25Retriever.from_texts(texts, metadatas=metadatas)
        retriever.k = 15  # Get more results for better filtering
        
        logger.info("Successfully built enhanced BM25 retriever with context awareness")
        return retriever
        
    except Exception as e:
        logger.error(f"Error building retriever: {str(e)}")
        raise ValueError(f"Failed to build retriever: {str(e)}")

# -----------------------------
# 4. Context-Aware Retrieval
# -----------------------------
def context_aware_retrieve(query: str, retriever: BM25Retriever, section_type: str, k: int = 10) -> List[Tuple[Any, float]]:
    """Retrieve with section-type awareness to avoid context confusion."""
    try:
        # Map ICF sections to relevant chunk types
        section_mapping = {
            "Purpose of the Study": ["purpose_info", "general"],
            "Study Procedures": ["planned_action", "general"],
            "Risks": ["risk_info", "general"], 
            "Benefits": ["benefit_info", "general"]
        }
        
        all_docs = retriever.get_relevant_documents(query)
        
        # Filter by relevant chunk types for this section
        relevant_types = section_mapping.get(section_type, ["general"])
        filtered_docs = [doc for doc in all_docs 
                        if doc.metadata.get("chunk_type", "general") in relevant_types]
        
        # Return with dummy scores for compatibility
        return [(doc, 1.0) for doc in filtered_docs[:k]]
        
    except Exception as e:
        logger.error(f"Error in context_aware_retrieve: {str(e)}")
        # Fallback to basic retrieval
        docs = retriever.get_relevant_documents(query)
        return [(doc, 1.0) for doc in docs[:k]]

# -----------------------------
# 5. ICF Generation with Context-Aware Retrieval
# -----------------------------
def generate_full_icf(retriever: BM25Retriever) -> Tuple[Dict[str, str], List[Tuple[Any, float]]]:
    """Generate ICF sections with context-aware retrieval."""
    try:
        all_results = []
        sections = {}
        
        # Generate each section with targeted retrieval
        section_queries = {
            "Purpose of the Study": "study purpose objective primary endpoint hypothesis",
            "Study Procedures": "procedures visits assessments treatment administration",
            "Risks": "adverse events side effects risks toxicity safety",
            "Benefits": "benefits potential improvement outcomes efficacy"
        }
        
        for section_name, query in section_queries.items():
            # Use context-aware retrieval for each section
            results = context_aware_retrieve(query, retriever, section_name, k=8)
            all_results.extend(results)
            
            # Build context for this specific section
            context = "\n".join([doc.page_content for doc, _ in results])
            
            if not context.strip():
                sections[section_name] = f"Information for {section_name} not found in protocol."
                continue
                
            # Generate section-specific content
            section_content = generate_section_content(section_name, context)
            sections[section_name] = section_content

        logger.info("Successfully generated ICF sections with context-aware retrieval")
        return sections, all_results

    except Exception as e:
        logger.error(f"Error generating ICF: {str(e)}")
        # Return fallback sections
        fallback_sections = {
            "Purpose of the Study": "Unable to extract purpose from protocol.",
            "Study Procedures": "Unable to extract procedures from protocol.",
            "Risks": "Unable to extract risks from protocol.",
            "Benefits": "Unable to extract benefits from protocol."
        }
        return fallback_sections, []

def generate_section_content(section_name: str, context: str) -> str:
    """Generate content for a specific ICF section."""
    
    # Customize prompts based on section
    if section_name == "Purpose of the Study":
        specific_instructions = "Focus only on the study's objectives, goals, and what the research aims to discover. Do not include patient numbers, duration, or procedural details."
    elif section_name == "Study Procedures":
        specific_instructions = "Include all procedural details, assessments, visits, timeline, number of patients, and study duration.It its essential to state the number of patients and study duration. Be specific about what participants will experience."
    elif section_name == "Risks":
        specific_instructions = "Focus only on potential adverse events, side effects, and safety concerns."
    elif section_name == "Benefits":
        specific_instructions = "Focus only on potential positive outcomes and benefits to participants."
    else:
        specific_instructions = "Be concise and clear."
    
    icf_prompt = f"""
You are drafting an Informed Consent Form (ICF) for a clinical trial participant.

Context from the clinical trial protocol:
{context}

Generate ONLY the "{section_name}" section of the ICF.

Rules:
- Use only information from the context.
- Write in plain 8th-grade language.
- Be concise and clear.
- {specific_instructions}
- Do not invent information.
- Return only the content for this section, no JSON formatting.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a careful clinical trial form writer."},
                {"role": "user", "content": icf_prompt}
            ],
            temperature=0.3
        )

        return response.choices[0].message.content.strip()
        
    except Exception as e:
        logger.error(f"Error generating section {section_name}: {str(e)}")
        return f"Unable to generate {section_name} from protocol."

# -----------------------------
# 6. Verification RAG Against Error Handling
# -----------------------------
def verify_output(draft_text: str, retriever: BM25Retriever, target_section: str) -> Dict[str, Any]:
    """Verify output against source material."""
    try:
        sentences = sent_tokenize(draft_text)
        verified = {}
        for sent in sentences:
            if sent.strip():
                retrieved = context_aware_retrieve(sent, retriever, target_section, k=3)
                support = [doc.page_content for doc, score in retrieved if score > 0.5]
                verified[sent] = support if support else "No strong support"
        return verified
    except Exception as e:
        logger.error(f"Error in verification: {str(e)}")
        return {"Error": f"Verification failed: {str(e)}"}

# -----------------------------
# 7. Judge LLM with Error Handling
# -----------------------------
def judge_output(section_name: str, draft_text: str, retriever: BM25Retriever) -> str:
    """Judge the quality and accuracy of generated sections."""
    try:
        evidence_docs = retriever.get_relevant_documents(draft_text)
        evidence = "\n\n".join([doc.page_content for doc in evidence_docs[:5]])

        judge_prompt = f"""
You are a clinical trial protocol review assistant.
Your task is to judge whether the following draft section of an Informed Consent Form (ICF)
is accurate and faithful to the supporting evidence from the trial protocol.

Section: {section_name}

Draft ICF Text:
{draft_text}

Retrieved Evidence:
{evidence}

Guidelines:
- Mark as Credible if the draft is fully supported by the evidence.
- Mark as Not Supported if claims are contradicted or missing in the evidence.
- Mark as Ambiguous if the evidence is incomplete or partially supports the draft.
- Be strict about numbers (patients, duration) and risks/benefits claims.
- Provide a short justification.

Return your answer in JSON with keys: decision, justification.
"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a careful and strict clinical trial protocol judge."},
                {"role": "user", "content": judge_prompt}
            ],
            temperature=0.1
        )

        return response.choices[0].message.content

    except Exception as e:
        logger.error(f"Error in judge_output: {str(e)}")
        return json.dumps({
            "decision": "Error", 
            "justification": f"Judgment failed: {str(e)}"
        })

# -----------------------------
# 8. Template Filling with Fallback
# -----------------------------
def fill_template(template_path: str, output_path: str, icf_sections: Dict[str, Dict[str, Any]]) -> None:
    """Fill template with generated sections by reading, modifying, and recreating."""
    try:
        if os.path.exists(template_path):
            # Read the template
            doc = Document(template_path)
            
            # Create a new document with the same content but filled sections
            new_doc = Document()
            
            section_mapping = {
                "Section 1. Purpose of the Research": "Purpose of the Study",
                "Section 2. Procedures": "Study Procedures", 
                "Section 4. Discomforts and Risks": "Risks",
                "Section 5. Potential Benefits": "Benefits"
            }
            
            # Copy all content, replacing target sections
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Check if this is one of our target sections
                section_to_fill = None
                for template_heading, gen_key in section_mapping.items():
                    if text == template_heading:
                        section_to_fill = gen_key
                        break
                
                # Add the original paragraph
                new_para = new_doc.add_paragraph(paragraph.text)
                
                # If this was a target section, add the generated content
                if section_to_fill and section_to_fill in icf_sections and "text" in icf_sections[section_to_fill]:
                    new_doc.add_paragraph(icf_sections[section_to_fill]["text"])
            
            new_doc.save(output_path)
        
        else:
            # Create basic document fallback
            doc = Document()
            doc.add_heading("Informed Consent Form", 0)
            
            for section_name in ["Purpose of the Study", "Study Procedures", "Risks", "Benefits"]:
                doc.add_heading(section_name, 1)
                if section_name in icf_sections and "text" in icf_sections[section_name]:
                    doc.add_paragraph(icf_sections[section_name]["text"])
                else:
                    doc.add_paragraph(f"[{section_name} information not available]")
            
            doc.save(output_path)

        logger.info(f"Successfully saved ICF document to: {output_path}")

    except Exception as e:
        logger.error(f"Error filling template: {str(e)}")
        raise ValueError(f"Failed to create ICF document: {str(e)}")
# -----------------------------
# 9. Enhanced Logging Export
# -----------------------------
def export_logs(sections: Dict[str, Dict[str, Any]], filename: str = "logs.json") -> None:
    """Export detailed logs with citations and metadata."""
    try:
        # Add timestamp and metadata to logs
        enhanced_logs = {
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "retrieval_strategy": "Enhanced context-aware BM25 with semantic chunk classification",
            "sections": sections,
            "metadata": {
                "total_sections": len(sections),
                "sections_with_citations": sum(1 for s in sections.values() if s.get("citations")),
                "agentic_features": [
                    "Context-aware chunk classification",
                    "Section-specific retrieval filtering", 
                    "Planned vs completed action distinction",
                    "Domain-aware metadata enrichment"
                ]
            }
        }
        
        with open(filename, "w") as f:
            json.dump(enhanced_logs, f, indent=2, default=str)
        
        logger.info(f"Successfully exported enhanced logs to: {filename}")
        
    except Exception as e:
        logger.error(f"Error exporting logs: {str(e)}")
        # Try to save basic version
        try:
            with open(filename, "w") as f:
                json.dump({"error": str(e), "sections": sections}, f, indent=2, default=str)
        except:
            logger.error("Failed to save even basic logs")
